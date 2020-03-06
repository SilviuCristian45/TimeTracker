import time
from docx import Document
import os
import datetime 
from datetime import date

running_state = "true"
today = str(date.today())
print("Today is : "+today)
start_hour = datetime.datetime.now().hour
start_minute = datetime.datetime.now().minute
print("You opened the PC : " + str(start_hour) +":"+ str(start_minute))
#running when the script starts
#1) take the most recent word document

documents = os.listdir('docs') # * means all if need specific format then *.csv
most_recent_document = documents[len(documents)-1]
most_recent_document = most_recent_document[:-5]

#2) if its another day , create another document
if (most_recent_document != today):
    document = Document()
    document.add_paragraph(
        str(start_hour) +":"+str(start_minute), style='List Bullet'
    )
    document.save("docs/"+str(today)+".docx")
else :
    #3) else just write in the most recent document
    document = Document("docs/"+str(today)+".docx")
    document.add_paragraph(
        str(start_hour) +":"+str(start_minute), style='List Bullet'
    )
    document.save("docs/"+str(today)+".docx")
    

#running while the script is opened
while True:
    running_state = input("Type false in order to close the program ...")
    #running when the script is closed
    if running_state == "false":
        document = Document("docs/"+str(today)+".docx")
        paragraphs = document.paragraphs

        start_hour = datetime.datetime.now().hour
        start_minute = datetime.datetime.now().minute

        paragraphs[len(paragraphs)-1].text = paragraphs[len(paragraphs)-1].text +"-"+str(start_hour) +":"+str(start_minute)
        document.save("docs/"+str(today)+".docx")
        break

print("-------------   CLOSE PROGRAM   ---------------")